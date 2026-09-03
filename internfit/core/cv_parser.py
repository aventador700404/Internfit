from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from io import BytesIO
from functools import lru_cache
import shutil
import subprocess
from typing import BinaryIO
from typing import Iterable
import re

from docx import Document
from pypdf import PdfReader
try:
    import pymupdf as fitz
except ImportError:  # pragma: no cover - production dependencies include PyMuPDF
    try:
        import fitz
    except ImportError:  # pragma: no cover - production dependencies include PyMuPDF
        fitz = None

try:
    from PIL import Image, ImageOps
except ImportError:  # pragma: no cover - Pillow is installed in production
    Image = None
    ImageOps = None

try:
    from pdfminer.high_level import extract_text as _extract_pdfminer_text
except ImportError:  # pragma: no cover - production dependencies include pdfminer.six
    _extract_pdfminer_text = None


EVIDENCE_RULES = {
    "strategy": (
        "strategy", "strategic", "portfolio", "bcg matrix", "five forces", "vrio",
        "전략", "전략기획", "사업전략", "경영전략", "기획",
    ),
    "research": (
        "research", "screened", "market", "consumer", "interview", "insight",
        "리서치", "시장조사", "시장 분석", "시장분석", "소비자 조사", "인사이트",
    ),
    "operations": (
        "operations", "operation", "planning", "execution", "coordination", "logistics",
        "운영", "운영관리", "프로세스 개선", "업무 개선", "실행", "조율",
    ),
    "stakeholder": (
        "stakeholder", "partnered", "communication", "community", "cross-cultural",
        "이해관계자", "유관부서", "파트너", "커뮤니케이션", "협업", "커뮤니티",
    ),
    "data_analysis": (
        "data", "analytics", "analysis", "quantitative", "excel", "financial",
        "데이터", "데이터 분석", "데이터분석", "통계", "정량분석", "엑셀",
    ),
    "technology": (
        "technology", "information systems", "digital", "platform", "ai", "generative ai",
        "기술", "디지털", "인공지능", "생성형 ai", "플랫폼", "시스템", "자동화",
    ),
    "finance": (
        "finance", "fintech", "investment", "financial", "m&a", "acquisition",
        "금융", "재무", "투자", "자본시장", "채권", "인수합병",
    ),
    "event_management": (
        "event", "recruitment", "onboarding", "training", "member",
        "행사", "이벤트", "워크숍", "교육 운영", "행사기획",
    ),
    "marketing": (
        "marketing", "social media", "campaign", "influencer", "content strategy", "brand",
        "마케팅", "소셜미디어", "캠페인", "인플루언서", "콘텐츠", "브랜딩",
    ),
    "sales": (
        "sales", "business development", "lead generation", "account management",
        "영업", "사업개발", "리드 발굴", "고객관리",
    ),
    # Keep generic credentials such as "SQL Developer" from turning a business
    # candidate into a software engineer.  A real engineering signal should be
    # an engineering role, a programming deliverable, or a concrete stack.
    "software_engineering": (
        "software engineer", "software engineering", "backend", "frontend", "api",
        "programming language", "programming project", "software developer", "coding",
        "소프트웨어 개발", "백엔드", "프론트엔드", "애플리케이션 개발", "프로그래밍",
        "코딩", "개발 프로젝트",
    ),
    "mcp_integration": (
        "mcp", "model context protocol", "rest api", "api integration", "mcp server", "mcp client",
        "모델 컨텍스트 프로토콜", "api 연동", "시스템 연동",
    ),
    "accounting": (
        "accounting", "audit", "journal entries", "reconciliation", "monthly close", "bookkeeping",
        "회계", "감사", "분개", "조정", "월말 마감", "장부",
    ),
}

LANGUAGE_RULES = {
    "korean": ("korean (native)", "korean", "한국어", "국어", "모국어"),
    "english": ("english (cefr c1", "toefl", "english", "영어", "토익", "토플", "토스", "오픽", "teps"),
    "german": ("german (cefr a2", "german", "독일어"),
    "japanese": ("japanese", "일본어", "jlpt"),
    "chinese": ("chinese", "mandarin", "중국어", "만다린", "hsk"),
}

TOOL_RULES = {
    "excel": ("excel", "엑셀"),
    "powerpoint": ("powerpoint", "파워포인트", "ppt", "피피티"),
    "word": ("word", "워드"),
    "sql": ("sqld", "sql", "에스큐엘"),
    "ai_tools": ("generative ai", "ai tools", "rapid prototyping", "생성형 ai", "생성형 인공지능", "프롬프트", "llm"),
    "python": ("python", "파이썬"),
    "java": ("java",),
    "javascript": ("javascript", "typescript", "자바스크립트", "타입스크립트"),
    "git": ("git", "github"),
    "docker": ("docker", "도커"),
    "kubernetes": ("kubernetes", "쿠버네티스"),
    "power_bi": ("power bi", "파워 bi"),
    "notion": ("notion", "노션"),
    "asana": ("asana", "아사나"),
    "sap": ("sap",),
    "erp": ("erp",),
    "figma": ("figma",),
}


@dataclass
class CandidateProfile:
    source_name: str
    raw_text: str
    evidence: dict[str, list[str]]
    languages: set[str]
    tools: set[str]
    graduation: str | None
    education: list[str]

    @property
    def evidence_tags(self) -> set[str]:
        return {tag for tag, lines in self.evidence.items() if lines}


def extract_docx_text(file_path: str | Path) -> list[str]:
    document = Document(file_path)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return lines


def extract_docx_bytes(data: bytes | BinaryIO) -> list[str]:
    """Extract visible DOCX text without persisting the uploaded file."""
    document = Document(BytesIO(data) if isinstance(data, bytes) else data)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return lines


def extract_pdf_bytes(data: bytes) -> list[str]:
    """Extract selectable PDF text without persisting the uploaded file."""
    try:
        reader = PdfReader(BytesIO(data))
        lines: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            lines.extend(line.strip() for line in text.splitlines() if line.strip())
        pypdf_text = "\n".join(lines)
        candidates = [pypdf_text]
        # Some Korean PDFs have a valid text layer but a font encoding that
        # pypdf cannot map back to Unicode. Try pdfminer in that case; this is
        # still text extraction, not OCR, and the uploaded bytes stay in memory.
        if _extract_pdfminer_text is not None and (
            not _pdf_text_is_usable(pypdf_text) or _hangul_count(pypdf_text) == 0
        ):
            try:
                candidates.append(_extract_pdfminer_text(BytesIO(data)) or "")
            except Exception:
                pass
        if not _pdf_text_is_usable(max(candidates, key=_pdf_text_quality)):
            candidates.append(_extract_pdf_with_ocr(data))
        best_text = max(candidates, key=_pdf_text_quality)
        if not _pdf_text_is_usable(best_text):
            return []
        return [line.strip() for line in best_text.splitlines() if line.strip()]
    except Exception as exc:
        raise ValueError("Could not read this PDF. Please upload a text-based PDF.") from exc


def _hangul_count(text: str) -> int:
    return sum("가" <= character <= "힣" for character in text)


def _pdf_text_quality(text: str) -> int:
    visible = [character for character in text if not character.isspace()]
    replacements = text.count("�")
    controls = sum(character.isprintable() is False for character in visible)
    return len(visible) + (_hangul_count(text) * 3) - (replacements * 20) - (controls * 10)


def _pdf_text_is_usable(text: str) -> bool:
    visible = [character for character in text if not character.isspace()]
    alphanumeric = sum(character.isalnum() for character in visible)
    replacements = text.count("�")
    return len(visible) >= 8 and alphanumeric >= 4 and replacements <= max(2, len(visible) // 100)


@lru_cache(maxsize=1)
def _tesseract_languages(executable: str) -> frozenset[str]:
    """Return installed Tesseract languages without exposing document text."""
    try:
        result = subprocess.run(
            [executable, "--list-langs"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
            timeout=10,
        )
    except (OSError, subprocess.SubprocessError):
        return frozenset()

    languages = set()
    for line in result.stdout.decode("utf-8", errors="replace").splitlines():
        value = line.strip()
        if re.fullmatch(r"[A-Za-z0-9_]+", value):
            languages.add(value)
    return frozenset(languages)


def _ocr_language_candidates(executable: str) -> list[str]:
    available = _tesseract_languages(executable)
    candidates: list[str] = []
    if {"kor", "eng"}.issubset(available):
        candidates.append("kor+eng")
    for language in ("kor", "eng"):
        if language in available and language not in candidates:
            candidates.append(language)
    # Keep a useful fallback if the language probe itself was unavailable.
    return candidates or ["kor+eng", "eng"]


def _ocr_png_variants(pixmap) -> list[tuple[str, bytes]]:
    """Create OCR-friendly variants while keeping all bytes in memory."""
    original = pixmap.tobytes("png")
    variants = [("original", original)]
    if Image is None or ImageOps is None:
        return variants

    try:
        with Image.open(BytesIO(original)) as source:
            grayscale = ImageOps.autocontrast(source.convert("L"))
            gray_buffer = BytesIO()
            grayscale.save(gray_buffer, format="PNG", optimize=True)
            variants.append(("gray", gray_buffer.getvalue()))

            # Colored resume sidebars can hide white/yellow text from OCR.
            # A high-contrast pass turns both the sidebar and white page into
            # readable black/white text without saving an intermediate file.
            binary = grayscale.point(lambda value: 255 if value >= 165 else 0)
            binary_buffer = BytesIO()
            binary.save(binary_buffer, format="PNG", optimize=True)
            variants.append(("binary", binary_buffer.getvalue()))
    except (OSError, ValueError):
        pass
    return variants


def _run_tesseract(
    executable: str,
    image_bytes: bytes,
    language: str,
    psm: str,
    page_number: int,
    variant: str,
) -> str:
    try:
        result = subprocess.run(
            [
                executable,
                "stdin",
                "stdout",
                "-l",
                language,
                "--oem",
                "1",
                "--psm",
                psm,
                "--dpi",
                "300",
                "-c",
                "preserve_interword_spaces=1",
            ],
            input=image_bytes,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
            timeout=10,
        )
    except subprocess.TimeoutExpired:
        print(
            f"[InternFit] OCR timeout page={page_number} language={language} "
            f"variant={variant} psm={psm}",
            flush=True,
        )
        return ""
    except (OSError, subprocess.SubprocessError) as exc:
        print(
            f"[InternFit] OCR process error page={page_number} language={language} "
            f"variant={variant} psm={psm} error={type(exc).__name__}",
            flush=True,
        )
        return ""

    text = result.stdout.decode("utf-8", errors="replace")
    if result.returncode != 0 or not text.strip():
        print(
            f"[InternFit] OCR attempt page={page_number} language={language} "
            f"variant={variant} psm={psm} rc={result.returncode} chars={len(text)}",
            flush=True,
        )
    return text


def _extract_pdf_with_ocr(data: bytes) -> str:
    """OCR image-only PDFs with Korean/English fallback passes.

    Pages are rendered and piped directly to Tesseract; no uploaded PDF or
    rendered page is persisted. A small number of layout and contrast passes
    are used because resumes often contain two columns, colored sidebars, or
    outlined text rather than a selectable text layer. The pass count is
    deliberately bounded so a free web instance does not spend minutes on
    one upload.
    """
    if fitz is None:
        return ""
    executable = shutil.which("tesseract")
    if not executable:
        print("[InternFit] OCR unavailable: tesseract executable not found", flush=True)
        return ""

    languages = _ocr_language_candidates(executable)
    try:
        document = fitz.open(stream=data, filetype="pdf")
        page_text: list[str] = []
        for page_number, page in enumerate(document, start=1):
            # Keep the previous 2x size for predictable latency on the free
            # instance; the contrast pass handles the problematic sidebar.
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            variants = dict(_ocr_png_variants(pixmap))
            primary = languages[0]
            attempts = (
                ("original", "3"),
                ("binary", "11"),
            )
            best_page = ""
            for variant, psm in attempts:
                image_bytes = variants.get(variant, variants["original"])
                text = _run_tesseract(
                    executable, image_bytes, primary, psm, page_number, variant
                )
                if _pdf_text_quality(text) > _pdf_text_quality(best_page):
                    best_page = text
                if _pdf_text_is_usable(best_page):
                    break

            # If the combined Korean+English model fails or returns no useful
            # text, try the individual installed language models as a safety
            # net. Keep this to one final attempt to bound request latency;
            # English-only text is still preferable to rejecting a CV.
            if not _pdf_text_is_usable(best_page):
                fallback = next(
                    (language for language in ("eng", "kor") if language in languages[1:]),
                    None,
                )
                if fallback:
                    image_bytes = variants.get("binary", variants["original"])
                    fallback_text = _run_tesseract(
                        executable, image_bytes, fallback, "11", page_number, "binary"
                    )
                    if _pdf_text_quality(fallback_text) > _pdf_text_quality(best_page):
                        best_page = fallback_text

            if best_page.strip():
                page_text.append(best_page)
        document.close()
        return "\n".join(page_text)
    except Exception as exc:
        print(f"[InternFit] OCR document error type={type(exc).__name__}", flush=True)
        return ""


def parse_pdf_bytes(data: bytes, source_name: str = "uploaded_cv.pdf") -> CandidateProfile:
    lines = extract_pdf_bytes(data)
    if not lines:
        raise ValueError(
            "No selectable text found in this PDF. Please upload a text-based PDF instead of a scanned image."
        )
    return _profile_from_lines(lines, source_name)


def _matching_lines(lines: Iterable[str], needles: tuple[str, ...]) -> list[str]:
    return [line for line in lines if any(_contains_term(line, needle) for needle in needles)]


def _contains_term(text: str, term: str) -> bool:
    normalized_text = re.sub(r"\s+", " ", text.casefold()).strip()
    normalized_term = re.sub(r"\s+", " ", term.casefold()).strip()
    if re.search(r"[가-힣]", normalized_term):
        # Korean particles and spacing vary in resumes, e.g. "시장 분석을"
        # versus "시장분석". Keep aliases specific enough to avoid broad
        # single-syllable matches, then compare a compact form.
        return re.sub(r"\s+", "", normalized_term) in re.sub(r"\s+", "", normalized_text)
    escaped = re.escape(normalized_term).replace(r"\ ", r"\s+")
    return bool(re.search(rf"(?<![a-z0-9]){escaped}(?![a-z0-9])", normalized_text))


def parse_cv(file_path: str | Path) -> CandidateProfile:
    lines = extract_docx_text(file_path)
    return _profile_from_lines(lines, Path(file_path).name)


def parse_cv_bytes(data: bytes, source_name: str = "uploaded_cv.docx") -> CandidateProfile:
    return _profile_from_lines(extract_docx_bytes(data), source_name)


def _profile_from_lines(lines: list[str], source_name: str) -> CandidateProfile:
    raw_text = "\n".join(lines)
    evidence = {tag: _matching_lines(lines, needles) for tag, needles in EVIDENCE_RULES.items()}
    languages = {
        language
        for language, needles in LANGUAGE_RULES.items()
        if any(_contains_term(raw_text, needle) for needle in needles)
    }
    tools = {
        tool
        for tool, needles in TOOL_RULES.items()
        if any(_contains_term(raw_text, needle) for needle in needles)
    }
    graduation_patterns = (
        r"\b(?:bachelor|master|ph\.?d|undergraduate|graduate)\s+(?:student|candidate)\b",
        r"\b(?:candidate|student)\b.*(?:\b20\d{2}\b|\buniversity\b|\bcollege\b)",
        r"\b(?:expected|anticipated)\s+(?:graduation|completion|to graduate)\b",
        r"\bclass of\s+20\d{2}\b",
        r"(?:졸업예정|재학)",
    )
    graduation = next(
        (
            line
            for line in lines
            if any(re.search(pattern, line, flags=re.IGNORECASE) for pattern in graduation_patterns)
        ),
        None,
    )
    education = [
        line
        for line in lines
        if re.search(r"\b(?:university|college|school|institute)\b", line, flags=re.IGNORECASE)
        or re.search(r"(?:대학교|대학|재학|학사|석사|박사|경영학|경제학|무역학|국제통상)", line)
